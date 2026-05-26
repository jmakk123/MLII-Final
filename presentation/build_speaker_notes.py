"""Build speaker_notes.docx - index card style with technical sub-explanations."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_speaker_header(name, time, pages):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1A, 0x4D, 0x8C)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"{time}   |   {pages}")
    r2.font.size = Pt(10)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_section(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(6)


def bullet(parts):
    p = doc.add_paragraph(style='List Bullet')
    for text, is_bold in parts:
        r = p.add_run(text)
        r.bold = is_bold
        r.font.size = Pt(11)


def sub(label, text):
    """Indented technical sub-explanation under a bullet."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.italic = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


# ============================================================
# Cover
# ============================================================
add_title("Speaker Notes")
add_subtitle("ML II Final / Forward 12-Month Max Drawdown / 12 Minutes")
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Bold = emphasize when speaking, or the UI element to point at.")
r.italic = True
r.font.size = Pt(10)
p = doc.add_paragraph()
r = p.add_run("Indented italic notes are technical explanations. Only say them if asked, or to pad time.")
r.italic = True
r.font.size = Pt(10)

# ============================================================
# Speaker 1
# ============================================================
doc.add_page_break()
add_speaker_header("Speaker 1 / Setup", "2:30", "Overview, Concepts, Data")

add_section("Open on the Overview page")
bullet([("What we predict: the ", False), ("biggest drop a stock takes in the next 12 months", True)])
sub("Why this number", "Drawdown is the asymmetric loss that hurts portfolios. Returns are mean-zero on average. Drawdowns are not. They define survival.")
bullet([("Why it matters: ", False), ("banks", True), (" use it for credit risk, ", False), ("funds", True), (" use it to size hedges, ", False), ("VCs", True), (" use it to gate follow-ons", False)])
bullet([("This site is the product, not a slide deck", False)])

add_section("Walk the four intro cards")
bullet([("Card 01: the problem. We do not know which firms fall hard next year", False)])
bullet([("Card 02: the data. ", False), ("20 years", True), (" of US public firms, ", False), ("about 30K firm-year snapshots", True)])
sub("What a snapshot is", "One firm at one fiscal year-end. Each snapshot is an independent training example. A firm appears in multiple years.")
bullet([("Card 03: the model. ", False), ("Small neural net", True), (" reads 5 years of accounting numbers plus recent price moves", False)])
bullet([("Card 04: the product. You can click any firm and see the forecast", False)])

add_section("Click into Concepts")
bullet([("Max drawdown in plain words: ", False), ("worst paper loss you sit through before the stock recovers", True)])
sub("Formal definition", "Over the next 252 trading days, find the running peak price. Drawdown is (current price minus running peak) divided by running peak. Max drawdown is the most negative value of that series.")
bullet([("Forward 12 months = looking ahead, not back. Model never sees the future in training", False)])
sub("How we enforce that", "Every input ratio and price feature is computed using data available at or before the snapshot date. The target is computed strictly after.")

add_section("Click into Data")
bullet([("Sources: ", False), ("CRSP", True), (" for prices, ", False), ("Compustat", True), (" for financials", False)])
sub("CRSP", "Center for Research in Security Prices. UChicago-housed academic database. Daily prices, returns, volumes for every US listed equity since 1925.")
sub("Compustat", "S&P-owned database of standardized fundamentals. Annual and quarterly income statement, balance sheet, cash flow items for US public firms.")
bullet([("Inputs: ", False), ("18 financial ratios", True), (" (Altman pieces, Ohlson, margins) plus ", False), ("7 price features", True), (" (volatility, momentum)", False)])
sub("The 18 ratios", "Altman X1 through X5 (working capital, retained earnings, EBIT, market cap, sales, each scaled by total assets or liabilities), Ohlson O-score components (size, leverage, profitability, liquidity), Zmijewski components (ROA, leverage, current ratio), and operating margins. All winsorized at 1 and 99 percent.")
sub("The 7 price features", "Trailing volatility, trailing momentum at multiple horizons, Sharpe ratio, trailing max drawdown, beta to the market. Computed off daily CRSP returns over the year before the snapshot.")
bullet([("Split: train ", False), ("2003-2017", True), (", validate ", False), ("2018-2019", True), (", test ", False), ("2020-2023", True)])
sub("Why time-blocked", "Random splits leak future information. A time-blocked split simulates the real deployment setting: predict next year using only past data.")
bullet([("Test period includes COVID, so the model had to survive a real shock", False)])

add_section("Hand off to Speaker 2")

# ============================================================
# Speaker 2
# ============================================================
doc.add_page_break()
add_speaker_header("Speaker 2 / Engine", "2:30", "Models, Findings")

add_section("Open the Models page")
bullet([("Model has ", False), ("two halves", True), (" that talk to each other", False)])
bullet([("Half 1 = ", False), ("small LSTM", True), (". Reads 18 ratios over 5 years, one year at a time", False)])
sub("LSTM", "Long Short-Term Memory network. A recurrent layer that processes a sequence step by step and keeps a hidden state that selectively remembers and forgets. Good at picking up trends and turning points across years.")
sub("Our config", "2 layers, hidden size 64, dropout 0.2. Input shape (5, 18). Output is the final hidden state, a 64-dim vector.")
bullet([("Half 2 = ", False), ("small feed-forward net", True), (". Reads recent price behavior", False)])
sub("MLP", "Multi-layer perceptron. Stacked dense layers with ReLU. Takes the 7-dim price feature vector and projects it into a 32-dim representation.")
bullet([("Both feed a ", False), ("fusion head", True), (" that outputs one number: forecast drawdown", False)])
sub("Fusion", "Concatenate the LSTM output (64) and MLP output (32) into a 96-dim vector. Pass through two dense layers. Final layer produces a single scalar bounded with tanh.")

add_section("Training, kept short")
bullet([("Loss: ", False), ("Huber", True), (" (forgiving of outliers) plus a small ", False), ("yes/no classifier head", True), (" asking \"drops more than 30%?\"", False)])
sub("Huber loss", "Behaves like squared error for small mistakes, like absolute error for large ones. The cutoff delta is 0.05. Stops one extreme firm from dominating the gradient.")
sub("The aux head", "A second output neuron with a binary cross-entropy loss on the label \"did this firm drop more than 30 percent.\" Weighted at 0.3. Forces the shared layers to also learn a discrimination signal, not just a regression signal.")
bullet([("The second head ", False), ("sharpened the ranking", True), (" without hurting the regression", False)])
bullet([("Trained ", False), ("3 seeds", True), (" and averaged. Ensemble beat any single run", False)])
sub("Ensembling", "Same architecture, same data, different random initializations. Average the three predictions at inference. Variance across seeds drops sharply and PR-AUC gains about 1.5 points.")
sub("Optimizer", "AdamW with weight decay 0.01, cosine annealing schedule, learning rate 1e-3, batch 256, early stopping with patience 8 on validation MAE.")

add_section("Switch to the Findings page")
bullet([("MAE: ", False), ("0.121", True), (". Forecast is off by about 12 percentage points on average", False)])
sub("MAE", "Mean Absolute Error. Average of the absolute difference between the predicted drawdown and the realized drawdown. Same units as the target.")
bullet([("PR-AUC: ", False), ("0.852", True), (". Deepest drops cluster near the top of the ranking", False)])
sub("PR-AUC", "Precision-Recall Area Under Curve, evaluated on the binary task \"did this firm drop more than 30 percent.\" Robust to class imbalance because deep drawdowns are rare.")
bullet([("Spearman: ", False), ("0.666", True), (". Predicted order matches the order that actually happens", False)])
sub("Spearman", "Rank correlation. Convert both predicted and actual values to ranks, then compute Pearson correlation. Measures whether the model ranks firms correctly, ignoring exact magnitudes.")

add_section("Honest caveat")
bullet([("Survivorship: ", False), ("firms that fully delist before the year ends drop out of the test set", True)])
sub("Why this matters", "A firm that goes bankrupt 3 months into the forward window has no full 12-month drawdown to compute. We mark these as delisted at -1.0 where flagged in CRSP (codes 02, 03, 04). For firms with incomplete forward windows the example is dropped. Drops the hardest cases out of the test set.")
bullet([("Test set is a little kinder than reality. We flag this on the page", False)])

add_section("Hand off to Speaker 3")

# ============================================================
# Speaker 3
# ============================================================
doc.add_page_break()
add_speaker_header("Speaker 3 / Product", "3:30", "Predictions, Top Risks, Compare, Backtest")

add_section("Live demo. Drive the site.")

add_section("Predictions page")
bullet([("Every prediction the model made on test years. ", False), ("Fifteen thousand of them", True)])
sub("What you are seeing", "15,311 rows. Each row is one firm at one snapshot date in 2020-2023. Columns: predicted drawdown, realized drawdown, outcome (hit, miss, safe, false alarm), sector.")
bullet([("Filter by ", False), ("year 2020", True), (" using year buttons", False)])
bullet([("Click a familiar firm row to open the ", False), ("firm detail drawer", True)])
bullet([("In the drawer, point at: ", False), ("5-year ratio history", True), (", ", False), ("forward 12-month price chart", True), (", ", False), ("predicted vs realized at top", True)])
sub("Ratio history", "The actual 5-year time series the LSTM consumed for this firm. Pulled from Compustat. Hover any line for the raw value at each fiscal year.")
sub("Price chart", "Daily closing price for the 252 trading days after the snapshot, normalized to 100. The dotted line shows the running peak. The shaded region is the drawdown path.")
bullet([("This is what a credit analyst sees. One screen, one decision", False)])

add_section("Top Risks page")
bullet([("Pick ", False), ("year 2022", True), (". Show the top 25 list", False)])
bullet([("Point at the ", False), ("hit rate badge", True), (" - percent of flagged firms that actually fell more than 30%", False)])
sub("Hit rate", "Of the top 25 firms ranked most risky by the model, count how many had a realized drawdown worse than -30 percent. Divide by 25. This is precision at top-K.")
bullet([("This is the watchlist a risk team opens at their morning meeting", False)])

add_section("Compare page (cut this first if running long)")
bullet([("We tested ", False), ("5 versions", True), (": financial LSTM alone, price alone, MLP, cross-attention, ", False), ("LSTM fusion ensemble", True)])
sub("Fin-only", "LSTM on the 18 ratios. No price features. Tests whether accounting fundamentals alone can rank drawdown risk.")
sub("Price-only", "MLP on the 7 price features. No fundamentals. Tests whether recent market behavior alone is enough.")
sub("MLP fusion", "Both feature sets, but the financial side is a flat MLP on the concatenated 5-year ratios, not an LSTM. Tests whether sequence modeling matters.")
sub("Cross-attention fusion", "Financial and price encoders attend to each other before fusing. More expressive but harder to train. Did not beat the simpler concatenation fusion.")
sub("LSTM fusion ensemble (winner)", "Our main architecture, averaged across 3 seeds. Best on all three metrics.")
bullet([("Fusion ensemble won. ", False), ("Reading both halves of a company beat reading either half alone", True)])

add_section("Backtest page")
bullet([("Walk across ", False), ("2020, 2021, 2022, 2023", True), (". Numbers do not collapse in any year", False)])
sub("What we plot", "MAE, PR-AUC, and hit rate computed within each test year separately. Stability across years is the actual deployment test.")
bullet([("Each year is a different kind of market: COVID, recovery, rate shock", False)])
bullet([("Model still ranks firms correctly across all of them", False)])

add_section("Hand off to Speaker 4")

# ============================================================
# Speaker 4
# ============================================================
doc.add_page_break()
add_speaker_header("Speaker 4 / Activity and Close", "3:30", "Activity, Use Cases")

add_section("Set up the game")
bullet([("Click ", False), ("Activity", True), (" in the sidebar", False)])
bullet([("Split the room into ", False), ("3 teams", True)])
bullet([("Each round: real firm at a real date. Vote yes or no on ", False), ("\"drops more than 30% in the next year?\"", True)])
sub("Why this framing", "Binary calls are easier to make under time pressure than point forecasts. They also map onto a real product use case: the watchlist filter.")
bullet([("Teams only see what the model saw. ", False), ("No future info", True)])
sub("What is shown", "Firm name, sector, snapshot date, the 5-year ratio history, and the trailing price chart. Same inputs the model used.")

add_section("Run the rounds")
bullet([("6 rounds, ", False), ("~25 seconds each", True)])
bullet([("Read firm name and snapshot date out loud", False)])
bullet([("Give 10 seconds to vote, then reveal", False)])
bullet([("Reveal: ", False), ("model prediction", True), (", then ", False), ("actual outcome", True)])
bullet([("1 point per correct call. Use on-screen scoring", False)])
bullet([("Final round overlay: ", False), ("FINAL ROUND ALL IN", True), (". Worth 2 points. Make it dramatic", False)])

add_section("After the game")
bullet([("Announce winning team", False)])
bullet([("Point at the ", False), ("model's score", True), (" on screen", False)])
bullet([("If model won: it does this for ", False), ("15,000 firms a year", True), (", not 6", False)])
bullet([("If team won: this is why we keep humans in the loop. Model is a co-pilot", False)])

add_section("Close on Use Cases")
bullet([("Click ", False), ("Use Cases", True), (" in the sidebar", False)])
bullet([("Banks: ", False), ("watchlist signals", True), (" for credit monitoring", False)])
sub("How it deploys", "Run the model monthly on every borrower in the loan book. Sort by drawdown score. Top decile gets routed to a human credit officer for review before any new exposure.")
bullet([("Funds: size ", False), ("hedges", True), (" off the score", False)])
sub("How it deploys", "For each long position, scale the put protection nominal as a function of the drawdown score. Higher score, deeper hedge. Lower score, lighter hedge. Cuts hedging cost without dropping protection where it matters.")
bullet([("VCs: gate ", False), ("follow-on capital", True), (" on worsening peer comps", False)])
sub("How it deploys", "For each portfolio company, identify the 5 closest public comparables. Track their drawdown scores monthly. If the comp basket deteriorates, slow or pause follow-on rounds in that sector.")
bullet([("Quants: ", False), ("long best decile, short worst", True), (" within sector", False)])
sub("How it deploys", "Rank firms within each sector. Go long the bottom decile (lowest predicted drawdown) and short the top decile (highest predicted drawdown). Rebalance quarterly. Sector-neutral by construction.")

add_section("Final line")
bullet([("Model is on the site. Data is on the site. ", False), ("Every prediction is on the site", True)])
bullet([("Thank you", False)])

# ============================================================
# Timing card
# ============================================================
doc.add_page_break()
add_section("Timing Card")
doc.add_paragraph()

table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Speaker"
hdr[1].text = "Section"
hdr[2].text = "Time"
hdr[3].text = "Stop At"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True

rows = [
    ("1", "Setup", "2:30", "2:30"),
    ("2", "Engine", "2:30", "5:00"),
    ("3", "Product", "3:30", "8:30"),
    ("4", "Activity and Close", "3:30", "12:00"),
    ("", "Total", "12:00", ""),
]
for i, (s, sec, t, stop) in enumerate(rows, 1):
    cells = table.rows[i].cells
    cells[0].text = s
    cells[1].text = sec
    cells[2].text = t
    cells[3].text = stop

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("If anyone runs over by more than 20 seconds, the next speaker should start anyway.")
r.italic = True
r.font.size = Pt(10)

from pathlib import Path
out_path = Path(__file__).resolve().parent / 'speaker_notes.docx'
doc.save(out_path)
print(f"Wrote {out_path}")
